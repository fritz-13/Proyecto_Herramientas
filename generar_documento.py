from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ══════════════════════════════════════════════════════
# CONFIGURACIÓN APA 7 (márgenes, fuente, interlineado)
# ══════════════════════════════════════════════════════
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ─── Estilos ───
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# Heading 1 style (APA Level 1: Centered, Bold)
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(12)
h1.font.bold = True
h1.font.italic = False
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(0)

# Heading 2 style (APA Level 2: Left, Bold)
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.italic = False
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(0)

# Heading 3 style (APA Level 3: Left, Bold Italic)
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(0)


# ─── Helper functions ───
def add_page_numbers():
    """Add page numbers to footer (right-aligned, Times New Roman)"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)


def add_heading_apa(text, level=1):
    """Add heading with APA formatting"""
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
    elif level == 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
    elif level == 3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
    return p


def add_para(text, bold_prefix="", indent=False):
    """Add paragraph with double spacing, APA style"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_code(text):
    """Add code block with single spacing, Consolas font"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p


def add_screenshot_placeholder(number, description):
    """Add screenshot placeholder in APA style"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"[CAPTURA DE PANTALLA {number}]")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run2 = p2.add_run(description)
    run2.font.italic = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run3 = p3.add_run("(Haga clic aquí e inserte la imagen)")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def add_bullet(text, bold_prefix=""):
    """Add bullet point with APA formatting"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def create_table(headers, rows):
    """Create table with APA styling"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════
# CARÁTULA (APA 7)
# ══════════════════════════════════════════════════════

# Title page content - centered, double spaced
title_lines = [
    ("", False),
    ("", False),
    ("", False),
    ("", False),
    ("UNIVERSIDAD TECNOLÓGICA DEL PERÚ", True),
    ("", False),
    ("Facultad de Ingeniería de Sistemas", False),
    ("", False),
    ("Herramientas de Desarrollo", False),
    ("", False),
    ("", False),
    ("AutoMarket Perú", True),
    ("Sistema de Seguimiento de Incidentes", False),
    ("con GitHub Issues, Integración Continua,", False),
    ("Entrega Continua y Plataformas en la Nube", False),
    ("", False),
    ("", False),
    ("Docente: Ing. Alex Santiago Sulca Onofre", False),
    ("", False),
    ("Ciclo: 2026 - I", False),
    ("", False),
    ("", False),
    ("Integrantes:", True),
]

for text, is_bold in title_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_bold:
        run.bold = True

# Student list in two columns on title page
students = [
    ("Bedriñana Huayta Denilson Fernando", "U23235023"),
    ("Campos Jimenez Juan Jesus", "U23248359"),
    ("Candiotti Alarcon Diego Gonzalo", "U22209741"),
    ("Carbajal Poma Jack Sheffer", "U20232976"),
    ("Castillon Torres Jhonny Alberto", "U23226379"),
    ("Ccolque Huachaca Hugo Andree", "U22249733"),
    ("Chavez Bruno Hualit Judeth", "U23205820"),
    ("Chumbe Sanchez Andre Emanuel", "U23251746"),
    ("Espinoza Flores Ginner", "U23220968"),
    ("Esteban Leiva Antony Raul", "U23226487"),
    ("Gaspar Chocce Jefferson Grenin", "U20226249"),
    ("Huaman Montes Renzo Brayan", "U22227724"),
    ("Huaman Salas Jack Jefferson", "U22300101"),
    ("Lopez Urco Jose Miguel", "U21322273"),
    ("Macha Capcha Jhyoan David", "U23226389"),
    ("Mallma Navarro Jarol Pier", "U23248487"),
    ("Melo Guillermo Jitsben Andree", "U23101151"),
    ("Mendoza Ricaldi Sofia Karol", "U23229905"),
    ("Meza Salazar Johann Rafael", "U23331396"),
    ("Mucha Filio Italo Fernando", "U23226271"),
    ("Peña Quispe Fritz Freddy", "U23300415"),
    ("Puchoc Castillo Jhoany Paoli", "U23208569"),
    ("Quintana Cajachagua Jhamir Ivan", "U22239944"),
    ("Quispe Mescua Eduardo Denis", "U20204450"),
    ("Ricse Elizalde Aaron Lenin", "U23247864"),
]

for name, code in students:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"{code} - {name}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

add_page_break()

# ══════════════════════════════════════════════════════
# ÍNDICE (Table of Contents)
# ══════════════════════════════════════════════════════
add_heading_apa("Índice", 1)

toc_items = [
    ("Introducción", 1),
    ("Marco Teórico", 1),
    ("Integración Continua (CI)", 2),
    ("Entrega Continua (CD)", 2),
    ("Plataformas en la Nube", 2),
    ("API RESTful", 2),
    ("GitHub Actions", 2),
    ("Lighthouse CI", 2),
    ("GitHub Issues", 2),
    ("Flask y Gunicorn", 2),
    ("Objetivos", 1),
    ("Objetivo General", 2),
    ("Objetivos Específicos", 2),
    ("Tecnologías Utilizadas", 1),
    ("Criterio 1: Sistema de Seguimiento de Incidentes (Semana 12)", 1),
    ("Criterio 2: Integración Continua - Lighthouse CI (Semana 13)", 1),
    ("Criterio 3: Entrega Continua - GitHub Pages (Semana 14)", 1),
    ("Criterio 4: Plataformas en la Nube (Semana 15)", 1),
    ("Conclusiones", 1),
    ("Recomendaciones", 1),
    ("Referencias", 1),
]

for item, level in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if level == 1:
        p.paragraph_format.left_indent = Cm(0)
        run = p.add_run(item)
        run.bold = True
    else:
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_page_break()

# ══════════════════════════════════════════════════════
# INTRODUCCIÓN
# ══════════════════════════════════════════════════════
add_heading_apa("Introducción", 1)
add_para(
    "El presente documento describe la implementación del proyecto AutoMarket Perú, "
    "una plataforma de venta de autos que integra un sistema de seguimiento de incidentes "
    "con herramientas modernas de desarrollo. El proyecto abarca cuatro criterios fundamentales: "
    "implementación de un sistema de seguimiento de incidentes con GitHub Issues, integración "
    "continua (CI) con Lighthouse CI, entrega continua (CD) mediante GitHub Actions, y "
    "despliegue en plataformas en la nube (GitHub Pages y Render).",
    indent=True
)
add_para(
    "AutoMarket Perú es una aplicación web estática que presenta un catálogo de vehículos en venta, "
    "con módulos de login, registro y un sistema de seguimiento de incidentes basado en GitHub Issues. "
    "GitHub Issues proporciona una plataforma nativa para la gestión de incidentes con labels, "
    "milestones, asignaciones y comentarios, sin necesidad de un backend adicional.",
    indent=True
)

# ══════════════════════════════════════════════════════
# MARCO TEÓRICO
# ══════════════════════════════════════════════════════
add_heading_apa("Marco Teórico", 1)

add_heading_apa("Integración Continua (CI)", 2)
add_para(
    "La Integración Continua es una práctica de desarrollo donde los miembros de un equipo "
    "integran su trabajo frecuentemente, generalmente varias veces al día. Cada integración "
    "se verifica mediante una compilación automatizada y pruebas, lo que permite detectar "
    "errores de manera temprana. En este proyecto, se implementó CI mediante Lighthouse CI, "
    "que ejecuta auditorías automáticas de calidad web en cada push al repositorio, "
    "verificando métricas de rendimiento, accesibilidad, buenas prácticas y SEO.",
    indent=True
)

add_heading_apa("Entrega Continua (CD)", 2)
add_para(
    "La Entrega Continua es una extensión de la integración continua que garantiza que "
    "el software pueda ser lanzado a producción en cualquier momento mediante despliegues "
    "automatizados. En este proyecto, se implementó CD con GitHub Actions, automatizando "
    "el despliegue del frontend estático a GitHub Pages cada vez que se realiza un push "
    "a las ramas principales del repositorio.",
    indent=True
)

add_heading_apa("Plataformas en la Nube", 2)
add_para(
    "Las plataformas en la nube ofrecen infraestructura como servicio (IaaS), plataforma "
    "como servicio (PaaS) y software como servicio (SaaS). En este proyecto se utilizaron "
    "GitHub Pages (hosting estático gratuito con CDN global) y Render (PaaS para el backend "
    "dinámico con Flask). Render proporciona HTTPS automático, escalado horizontal, logs en "
    "tiempo real y soporte nativo para aplicaciones Python con Gunicorn.",
    indent=True
)

add_heading_apa("API RESTful", 2)
add_para(
    "Una API RESTful (Representational State Transfer) es un estilo de arquitectura de "
    "software para sistemas distribuidos que utiliza los métodos HTTP estándar (GET, POST, "
    "PUT, DELETE) para realizar operaciones CRUD sobre recursos. En este proyecto, se "
    "desarrolló una API RESTful con Flask que expone endpoints para la gestión completa "
    "de incidentes, permitiendo la interacción entre el dashboard frontend y la base de "
    "datos SQLite.",
    indent=True
)

add_heading_apa("GitHub Actions", 2)
add_para(
    "GitHub Actions es una plataforma de automatización que permite crear pipelines de "
    "CI/CD directamente desde el repositorio de GitHub. Utiliza workflows definidos en "
    "archivos YAML que se ejecutan en respuesta a eventos como pushes, pull requests o "
    "schedules. Cada workflow está compuesto por uno o más jobs que se ejecutan en "
    "entornos virtuales (runners) y pueden incluir múltiples pasos para construir, probar "
    "y desplegar el software.",
    indent=True
)

add_heading_apa("Lighthouse CI", 2)
add_para(
    "Lighthouse CI es una herramienta de Google que automatiza la ejecución de auditorías "
    "de Lighthouse en pipelines de CI. Evalúa las aplicaciones web en cinco categorías: "
    "rendimiento (Performance), accesibilidad (Accessibility), mejores prácticas (Best "
    "Practices), SEO y Progressive Web App (PWA). Permite establecer umbrales mínimos de "
    "calidad que deben cumplirse para que el pipeline sea exitoso, garantizando así "
    "estándares de calidad en cada despliegue.",
    indent=True
)

add_heading_apa("GitHub Issues", 2)
add_para(
    "GitHub Issues es un sistema de seguimiento de incidencias integrado en GitHub "
    "que permite a los equipos de desarrollo gestionar bugs, mejoras y tareas directamente "
    "desde el repositorio. Cada issue puede tener labels (etiquetas) para categorización, "
    "milestones para agrupar objetivos, asignaciones a miembros del equipo y comentarios "
    "para discusión. Los issues se integran con commits y pull requests, proporcionando "
    "trazabilidad completa entre el código y las incidencias.",
    indent=True
)

add_heading_apa("Flask y Gunicorn", 2)
add_para(
    "Flask es un framework web ligero para Python que permite crear aplicaciones web y "
    "APIs RESTful de manera rápida y sencilla. Gunicorn (Green Unicorn) es un servidor "
    "WSGI HTTP para aplicaciones Python que actúa como interfaz entre el servidor web y "
    "la aplicación Flask, manejando múltiples peticiones concurrentes de manera eficiente. "
    "En producción, Gunicorn se utiliza como servidor de aplicación, reemplazando el "
    "servidor de desarrollo integrado de Flask.",
    indent=True
)

# ══════════════════════════════════════════════════════
# OBJETIVOS
# ══════════════════════════════════════════════════════
add_heading_apa("Objetivos", 1)
add_heading_apa("Objetivo General", 2)
add_para(
    "Implementar una plataforma web de venta de autos con un sistema de seguimiento de incidentes, "
    "aplicando herramientas de integración continua, entrega continua y despliegue en la nube.",
    indent=True
)
add_heading_apa("Objetivos Específicos", 2)
objectives = [
    "Configurar GitHub Issues como sistema de seguimiento de incidentes con labels personalizados y milestones.",
    "Crear y gestionar issues representando bugs, mejoras e incidentes del sistema.",
    "Configurar integración continua con Lighthouse CI para auditoría de calidad.",
    "Implementar entrega continua con GitHub Actions para deploy automático a GitHub Pages.",
    "Desplegar el frontend en GitHub Pages y el backend en Render.",
    "Documentar cada etapa del proceso con capturas de pantalla.",
]
for obj in objectives:
    add_bullet(obj)

# ══════════════════════════════════════════════════════
# TECNOLOGÍAS
# ══════════════════════════════════════════════════════
add_heading_apa("Tecnologías Utilizadas", 1)
add_para("A continuación se presentan las tecnologías empleadas en el desarrollo del proyecto:", indent=True)
tech_headers = ["Tecnología", "Versión", "Propósito"]
tech_rows = [
    ["HTML5", "5", "Estructura del frontend"],
    ["CSS3", "3", "Estilos y diseño responsive"],
    ["JavaScript", "ES6", "Interactividad del dashboard"],
    ["Python", "3.10.7", "Lenguaje del backend"],
    ["Flask", "3.1.3", "Framework web API REST"],
    ["SQLite", "3.x", "Base de datos"],
    ["Gunicorn", "26.0.0", "Servidor WSGI para producción"],
    ["GitHub Issues", "-", "Sistema de seguimiento de incidentes"],
    ["GitHub Actions", "-", "CI/CD automatizado"],
    ["Lighthouse CI", "0.14.x", "Auditoría de calidad"],
    ["GitHub Pages", "-", "Hosting frontend estático"],
    ["Render", "-", "Hosting backend (PaaS)"],
]
create_table(tech_headers, tech_rows)

add_page_break()

# ══════════════════════════════════════════════════════
# CRITERIO 1
# ══════════════════════════════════════════════════════
add_heading_apa("Criterio 1: Sistema de Seguimiento de Incidentes (Semana 12)", 1)
add_para("Puntaje: 5 puntos", bold_prefix="Nota: ")

add_heading_apa("Descripción", 2)
add_para(
    "Se implementó un sistema de seguimiento de incidentes utilizando GitHub Issues, "
    "una herramienta nativa de GitHub que permite gestionar bugs, mejoras y reportes "
    "del sistema mediante issues, labels y milestones. Este enfoque elimina la necesidad "
    "de un backend adicional para la gestión de incidentes, aprovechando la infraestructura "
    "de GitHub.",
    indent=True
)

add_heading_apa("Configuración de Labels Personalizados", 2)
add_para(
    "Se configuraron labels personalizados en el repositorio para categorizar los "
    "incidentes según tres dimensiones: tipo, prioridad y estado. Los labels permiten "
    "filtrar y organizar los issues de manera eficiente directamente desde GitHub:",
    indent=True
)
labels_headers = ["Dimensión", "Labels", "Propósito"]
labels_rows = [
    ["Tipo", "bug, mejora, incidente", "Clasifica la naturaleza del issue"],
    ["Prioridad", "prioridad: baja, prioridad: media, prioridad: alta, prioridad: crítica", "Indica el nivel de urgencia"],
    ["Estado", "estado: abierto, estado: progreso, estado: resuelto, estado: cerrado", "Indica la etapa del ciclo de vida"],
]
create_table(labels_headers, labels_rows)

add_heading_apa("Creación de Milestones", 2)
add_para(
    "Se creó un milestone titulado 'Semana 12 - Sistema de Seguimiento de Incidentes' "
    "con fecha de vencimiento al 17 de julio de 2026. Los milestones en GitHub Issues "
    "permiten agrupar issues relacionados y trackear el progreso hacia un objetivo "
    "específico, mostrando el porcentaje de issues completados vs. pendientes.",
    indent=True
)

add_heading_apa("Issues Creados (Incidentes)", 2)
add_para(
    "Se crearon 5 issues representando diferentes tipos de incidentes, cada uno con "
    "los labels correspondientes y asignados al milestone de la Semana 12. Los issues "
    "incluyen descripciones detalladas, pasos para reproducir (en caso de bugs), "
    "propuestas de solución y contexto del ambiente:",
    indent=True
)
issues_headers = ["#", "Título", "Labels", "Descripción"]
issues_rows = [
    ["1", "Error al cargar imágenes de autos en la página principal", "bug, prioridad: alta, estado: abierto", "Las imágenes del catálogo no se cargan, solo muestran placeholder"],
    ["2", "Agregar campo de búsqueda por marca de auto", "mejora, prioridad: media, estado: abierto", "Propuesta de filtro en tiempo real para el catálogo"],
    ["3", "Caída del servidor en horas pico - Error 502", "incidente, prioridad: crítica, estado: abierto", "Error 502 Bad Gateway entre 6-9 PM, posible sobrecarga"],
    ["4", "Formulario de registro no valida correos duplicados", "bug, prioridad: media, estado: abierto", "El registro sobrescribe cuentas con el mismo correo en localStorage"],
    ["5", "Implementar modo oscuro en el panel de administración", "mejora, prioridad: baja, estado: abierto", "Toggle de tema claro/oscuro con preferencia del sistema"],
]
create_table(issues_headers, issues_rows)

add_heading_apa("Flujo de Trabajo con GitHub Issues", 2)
add_para("El flujo de trabajo para el seguimiento de incidentes sigue estos pasos:", indent=True)
workflow_steps = [
    "Se crea un issue en GitHub describiendo el incidente con título, descripción detallada y ambiente.",
    "Se asignan labels de tipo (bug/mejora/incidente), prioridad (baja/media/alta/crítica) y estado (abierto).",
    "Se asigna el issue al milestone correspondiente para trackear el progreso.",
    "El equipo puede comentar en el issue para discutir soluciones, adjuntar capturas de pantalla y referenciar commits.",
    "Cuando se implementa una solución, se actualiza el estado a 'resuelto' y se referencia el PR o commit que lo soluciona.",
    "Una vez verificado, se cierra el issue y se mueve al estado 'cerrado'.",
]
for s in workflow_steps:
    add_bullet(s)

add_heading_apa("Ventajas de Usar GitHub Issues", 2)
adv_features = [
    "No requiere infraestructura adicional: está integrado directamente en GitHub.",
    "Soporte nativo para labels, milestones, asignaciones y comentarios.",
    "Integración con commits y pull requests para trazabilidad completa.",
    "Sistema de notificaciones por correo electrónico para cambios en issues.",
    "Búsqueda avanzada con filtros por label, milestone, asignado, fecha y más.",
    "API RESTful de GitHub Issues para integración con herramientas externas.",
    "Historial completo de cambios y discusiones en cada issue.",
]
for feat in adv_features:
    add_bullet(feat)

add_heading_apa("Acceso al Sistema", 2)
add_para(
    "El sistema de seguimiento de incidentes está accesible públicamente en la sección "
    "Issues del repositorio de GitHub:",
    indent=True
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p.add_run("https://github.com/fritz-13/Proyecto_Herramientas/issues")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

add_heading_apa("Estructura de Archivos del Proyecto", 2)
add_para("El proyecto está organizado con la siguiente estructura de directorios:", indent=True)
add_code("""Proyecto_Herramientas/
├── index.html               # Página principal con catálogo de autos
├── incidentes.html           # Dashboard de incidentes (visualización)
├── Login-register/
│   ├── login.html            # Página de inicio de sesión
│   └── register.html         # Página de registro
├── backend/
│   ├── app.py                # API Flask (incidentes vía GitHub Issues API)
│   ├── wsgi.py               # Punto de entrada WSGI
│   ├── __init__.py           # Marcador de paquete
│   ├── requirements.txt      # Dependencias Python
│   └── incidentes.db         # Base de datos SQLite
├── .github/workflows/
│   ├── deploy.yml            # CD - Deploy a GitHub Pages
│   └── lighthouse.yml        # CI - Auditoría Lighthouse
├── .lighthouserc.json         # Configuración de Lighthouse CI
└── render.yaml               # Configuración de Render""")

add_heading_apa("Evidencias", 2)
add_screenshot_placeholder(1, "Lista de issues en GitHub mostrando labels personalizados y milestone")
add_screenshot_placeholder(2, "Detalle de un issue con comentarios, labels y asignación")
add_screenshot_placeholder(3, "Vista de milestones mostrando el progreso de los issues")
add_screenshot_placeholder(4, "Labels personalizados configurados en el repositorio")

add_page_break()

# ══════════════════════════════════════════════════════
# CRITERIO 2
# ══════════════════════════════════════════════════════
add_heading_apa("Criterio 2: Integración Continua - Lighthouse CI (Semana 13)", 1)
add_para("Puntaje: 5 puntos", bold_prefix="Nota: ")

add_heading_apa("Descripción", 2)
add_para(
    "Se implementó un pipeline de Integración Continua (CI) utilizando Lighthouse CI "
    "a través de GitHub Actions. Este workflow se ejecuta automáticamente en cada push "
    "o pull request a la rama master, auditando la calidad del proyecto en términos de "
    "rendimiento, accesibilidad, buenas prácticas y SEO.",
    indent=True
)

add_heading_apa("Workflow de Lighthouse CI", 2)
add_para("Archivo: .github/workflows/lighthouse.yml", indent=True)
add_code("""name: Lighthouse CI Audit
on:
  push:
    branches: [master, main]
  pull_request:
    branches: [master, main]
jobs:
  lighthouse:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - name: Use Lighthouse CI
        run: |
          npm install -g @lhci/cli@0.14.x
          lhci autorun --config=.lighthouserc.json
        env:
          LHCI_GITHUB_APP_TOKEN: ${{ secrets.LHCI_GITHUB_APP_TOKEN }}""")

add_heading_apa("Configuración de Lighthouse CI", 2)
add_para("Archivo: .lighthouserc.json", indent=True)
add_code("""{
  "ci": {
    "collect": {
      "startServerCommand": "npx http-server . -p 8080",
      "url": [
        "http://localhost:8080/index.html",
        "http://localhost:8080/Login-register/login.html",
        "http://localhost:8080/Login-register/register.html",
        "http://localhost:8080/incidentes.html"
      ],
      "numberOfRuns": 1
    },
    "assert": {
      "assertions": {
        "categories:performance": ["warn", {"minScore": 0.6}],
        "categories:accessibility": ["warn", {"minScore": 0.7}],
        "categories:best-practices": ["warn", {"minScore": 0.7}],
        "categories:seo": ["warn", {"minScore": 0.7}]
      }
    },
    "upload": {
      "target": "temporary-public-storage"
    }
  }
}""")

add_heading_apa("Páginas Auditadas", 2)
pages_audited = [
    "index.html - Página principal con catálogo de autos.",
    "Login-register/login.html - Página de inicio de sesión.",
    "Login-register/register.html - Página de registro de usuarios.",
    "incidentes.html - Dashboard del sistema de incidentes.",
]
for p in pages_audited:
    add_bullet(p)

add_heading_apa("Métricas Evaluadas", 2)
add_para("Los umbrales de calidad configurados son los siguientes:", indent=True)
metrics_headers = ["Métrica", "Umbral Mínimo", "Descripción"]
metrics_rows = [
    ["Performance", "0.6 (60%)", "Velocidad de carga y rendimiento general"],
    ["Accessibility", "0.7 (70%)", "Accesibilidad para usuarios con discapacidades"],
    ["Best Practices", "0.7 (70%)", "Buenas prácticas de desarrollo web"],
    ["SEO", "0.7 (70%)", "Optimización para motores de búsqueda"],
]
create_table(metrics_headers, metrics_rows)

add_heading_apa("Análisis de Métricas de Lighthouse", 2)
add_para(
    "Lighthouse evalúa cada página web en múltiples categorías, generando puntuaciones "
    "del 0 al 100. A continuación se describen las categorías evaluadas y su importancia "
    "para el proyecto:",
    indent=True
)

lh_detail_headers = ["Categoría", "Rango de Puntaje", "Importancia para el Proyecto"]
lh_detail_rows = [
    ["Performance", "0-100", "Evalúa velocidad de carga, Core Web Vitals (LCP, FID, CLS) y optimización de recursos."],
    ["Accessibility", "0-100", "Verifica que la aplicación sea usable por personas con discapacidades (contraste, ARIA, navegación por teclado)."],
    ["Best Practices", "0-100", "Comprueba el uso de HTTPS, ausencia de vulnerabilidades, console errors y prácticas modernas."],
    ["SEO", "0-100", "Analiza meta tags, estructura semántica, robots.txt y optimización para motores de búsqueda."],
]
create_table(lh_detail_headers, lh_detail_rows)

add_para(
    "Los umbrales configurados en .lighthucerson.json establecen puntuaciones mínimas "
    "que deben alcanzarse para que el workflow de CI sea considerado exitoso. Si alguna "
    "página no cumple con los umbrales, el pipeline genera una advertencia pero no "
    "bloquea el despliegue, permitiendo a los desarrolladores identificar áreas de mejora.",
    indent=True
)

add_heading_apa("Beneficios de la Integración Continua con Lighthouse", 2)
ci_benefits = [
    "Detección temprana de regresiones de rendimiento antes de llegar a producción.",
    "Histórico de auditorías que permite trackear la evolución de la calidad del sitio.",
    "Reportes automáticos almacenados temporalmente para revisión del equipo.",
    "Integración directa con GitHub Actions sin necesidad de infraestructura adicional.",
    "Establece una línea base de calidad que todo nuevo código debe respetar.",
]
for b in ci_benefits:
    add_bullet(b)

add_heading_apa("Evidencias", 2)
add_screenshot_placeholder(5, "Workflow de Lighthouse CI ejecutándose en GitHub Actions")
add_screenshot_placeholder(6, "Resultados de la auditoría Lighthouse con puntuaciones")
add_screenshot_placeholder(7, "Reporte HTML generado por Lighthouse CI")

add_page_break()

# ══════════════════════════════════════════════════════
# CRITERIO 3
# ══════════════════════════════════════════════════════
add_heading_apa("Criterio 3: Entrega Continua - GitHub Pages (Semana 14)", 1)
add_para("Puntaje: 5 puntos", bold_prefix="Nota: ")

add_heading_apa("Descripción", 2)
add_para(
    "Se implementó un pipeline de Entrega Continua (CD) utilizando GitHub Actions "
    "para desplegar automáticamente el frontend estático en GitHub Pages cada vez que "
    "se realiza un push a la rama master o main.",
    indent=True
)

add_heading_apa("Workflow de Deploy", 2)
add_para("Archivo: .github/workflows/deploy.yml", indent=True)
add_code("""name: Deploy to GitHub Pages
on:
  push:
    branches: [master, main]
permissions:
  contents: read
  pages: write
  id-token: write
concurrency:
  group: pages
  cancel-in-progress: true
jobs:
  build:
    runs-on: ubuntu-latest
    steps:
      - name: Checkout
        uses: actions/checkout@v4
      - name: Setup Pages
        uses: actions/configure-pages@v5
      - name: Upload Pages artifact
        uses: actions/upload-pages-artifact@v3
        with:
          path: '.'
  deploy:
    needs: build
    runs-on: ubuntu-latest
    steps:
      - name: Deploy to GitHub Pages
        uses: actions/deploy-pages@v4""")

add_heading_apa("Flujo de Trabajo", 2)
add_para("El pipeline de CD funciona de la siguiente manera:", indent=True)
cd_steps = [
    "El desarrollador realiza un push a la rama master o main.",
    "GitHub Actions detecta el evento push e inicia el workflow.",
    "Job 'build': Clona el repositorio y prepara los archivos estáticos.",
    "Job 'build': Configura GitHub Pages con la acción configure-pages.",
    "Job 'build': Comprime y sube los archivos como artifact.",
    "Job 'deploy': Toma el artifact y lo despliega en GitHub Pages.",
    "El sitio se actualiza automáticamente en la URL pública.",
]
for i, step in enumerate(cd_steps, 1):
    add_bullet(f"Paso {i}: {step}")

add_heading_apa("URL del Frontend Desplegado", 2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
run = p.add_run("https://fritz-13.github.io/herramientas-de-desarrollo-practica")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

add_heading_apa("Verificación del Despliegue", 2)
add_para(
    "Para verificar que el despliegue se realizó correctamente, se siguieron los siguientes pasos:",
    indent=True
)
verify_steps = [
    "Acceder a la URL pública del frontend y verificar que todas las páginas carguen correctamente.",
    "Confirmar que los enlaces entre páginas funcionan (login, registro, catálogo, dashboard).",
    "Verificar que el dashboard de incidentes se visualiza correctamente desde GitHub Pages.",
    "Comprobar que la sección de Issues de GitHub es accesible desde el enlace en el dashboard.",
    "Revisar los logs de GitHub Actions para confirmar que el workflow se ejecutó sin errores.",
]
for v in verify_steps:
    add_bullet(v)

add_heading_apa("Evidencias", 2)
add_screenshot_placeholder(8, "Workflow de Deploy ejecutándose exitosamente en GitHub Actions")
add_screenshot_placeholder(9, "Página principal del sitio desplegado en GitHub Pages")
add_screenshot_placeholder(10, "Configuración de GitHub Pages en el repositorio")
add_screenshot_placeholder(11, "Historial de deploys exitosos")

add_page_break()

# ══════════════════════════════════════════════════════
# CRITERIO 4
# ══════════════════════════════════════════════════════
add_heading_apa("Criterio 4: Plataformas en la Nube (Semana 15)", 1)
add_para("Puntaje: 5 puntos", bold_prefix="Nota: ")

add_heading_apa("Descripción", 2)
add_para(
    "Se desplegó el proyecto en dos plataformas en la nube: GitHub Pages para el frontend "
    "estático y Render para el backend Flask. Esto permite que la aplicación sea accesible "
    "públicamente desde cualquier lugar con conexión a internet.",
    indent=True
)

add_heading_apa("Arquitectura Cloud", 2)
add_para("La arquitectura del proyecto en la nube sigue el siguiente esquema:", indent=True)
arch_items = [
    "El usuario accede mediante HTTPS al frontend alojado en GitHub Pages.",
    "El frontend (HTML/CSS/JS) se sirve desde la CDN global de GitHub Pages.",
    "El dashboard de incidentes realiza peticiones HTTP a la API del backend en Render.",
    "Render ejecuta el servidor Flask utilizando Gunicorn como WSGI server.",
    "Flask procesa las solicitudes y consulta la base de datos SQLite.",
]
for item in arch_items:
    add_bullet(item)

add_heading_apa("Configuración de Render (render.yaml)", 2)
add_code("""services:
  - type: web
    name: automarket-backend
    env: python
    buildCommand: pip install -r backend/requirements.txt
    startCommand: gunicorn backend.app:app
    envVars:
      - key: PYTHON_VERSION
        value: 3.10.7""")

add_heading_apa("Punto de Entrada WSGI (backend/wsgi.py)", 2)
add_code("""import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from app import app

if __name__ == "__main__":
    app.run()""")

add_heading_apa("Dependencias (requirements.txt)", 2)
add_code("""flask
flask-cors
gunicorn""")

add_heading_apa("URLs de los Servicios Desplegados", 2)
urls_headers = ["Servicio", "Plataforma", "URL"]
urls_rows = [
    ["Frontend", "GitHub Pages", "https://fritz-13.github.io/herramientas-de-desarrollo-practica"],
    ["API Backend", "Render", "https://herramientas-de-desarrollo-practica.onrender.com/api/incidentes"],
]
create_table(urls_headers, urls_rows)

add_heading_apa("Características de las Plataformas", 2)

add_heading_apa("GitHub Pages", 3)
gh_pages = [
    "Hosting gratuito para sitios web estáticos.",
    "CDN global con HTTPS automático y certificados SSL.",
    "Integración directa con GitHub Actions para CI/CD.",
    "Actualización automática del sitio con cada push al repositorio.",
    "Sin necesidad de configuración de servidor ni mantenimiento.",
]
for item in gh_pages:
    add_bullet(item)

add_heading_apa("Render", 3)
render_features = [
    "Plataforma como Servicio (PaaS) con capa gratuita.",
    "Soporte nativo para aplicaciones Python/Flask con Gunicorn.",
    "HTTPS automático con certificados SSL renovables.",
    "Despliegue automático desde GitHub con cada commit.",
    "Escalado horizontal automático según la demanda.",
    "Logs en tiempo real para monitoreo y debugging.",
    "Gestiona el ciclo de vida completo de la aplicación.",
]
for item in render_features:
    add_bullet(item)

add_heading_apa("Desafíos y Soluciones en el Despliegue", 2)
add_para("Durante el proceso de despliegue en la nube se enfrentaron los siguientes desafíos:", indent=True)

challenges_headers = ["Desafío", "Solución"]
challenges_rows = [
    ["ModuleNotFoundError: No module named 'backend.app' en Render. El archivo app.py no estaba en el repositorio remoto.",
     "Se agregó app.py, __init__.py y wsgi.py al repositorio mediante commit directo y se corrigió el startCommand a 'gunicorn backend.app:app'."],
    ["Render usa Python 3.14.3 por defecto, ignorando PYTHON_VERSION: 3.10.7 en render.yaml.",
     "Se verificó que la aplicación funciona correctamente con Python 3.14.3. No fue necesario forzar la versión."],
    ["El workflow de GitHub Pages fallaba por falta del environment 'github-pages'.",
     "Se eliminó el 'environment:' del workflow deploy.yml ya que no era necesario."],
    ["El repositorio original del grupo (JacksHS/Proyecto_Herramientas) no era accesible.",
     "Se creó un fork en fritz-13/herramientas-de-desarrollo-practica y se forzó el push."],
    ["El frontend mostraba error de CORS al conectar con la API en Render.",
     "Se habilitó Flask-CORS en app.py con CORS(app) para permitir peticiones desde cualquier origen."],
    ["Los Issues de GitHub estaban desactivados en el repositorio por ser un fork.",
     "Se habilitaron los Issues mediante PATCH a la API de GitHub (has_issues: true)."],
    ["El token fine-grained de GitHub no tenía permisos para crear labels e issues.",
     "Se creó un token clásico con scope public_repo y se usó la API REST de GitHub Issues."],
]
create_table(challenges_headers, challenges_rows)

add_heading_apa("Comparativa de Plataformas Cloud", 2)
add_para("La siguiente tabla compara las características de las plataformas utilizadas:", indent=True)
cloud_headers = ["Característica", "GitHub Pages", "Render"]
cloud_rows = [
    ["Tipo", "Hosting Estático", "PaaS (Platform as a Service)"],
    ["Backend", "No soportado", "Python, Node.js, Ruby, Go, etc."],
    ["Base de Datos", "No disponible", "SQLite, PostgreSQL, Redis"],
    ["HTTPS", "Automático (SSL)", "Automático (SSL)"],
    ["CDN", "Sí (Cloudflare)", "No"],
    ["Capa Gratuita", "Sí (ilimitado)", "Sí (limitado)"],
    ["Auto-deploy", "Mediante GitHub Actions", "Automático desde GitHub"],
    ["Logs", "No disponible", "En tiempo real"],
    ["Escalado", "Automático (CDN)", "Horizontal automático"],
]
create_table(cloud_headers, cloud_rows)

add_heading_apa("Evidencias", 2)
add_screenshot_placeholder(12, "Dashboard de Render mostrando el servicio activo y su estado")
add_screenshot_placeholder(13, "Logs de Render mostrando el build exitoso")
add_screenshot_placeholder(14, "API backend respondiendo correctamente desde la nube")
add_screenshot_placeholder(15, "Configuración de GitHub Pages en Settings del repositorio")

add_page_break()

# ══════════════════════════════════════════════════════
# CONCLUSIONES
# ══════════════════════════════════════════════════════
add_heading_apa("Conclusiones", 1)
conclusions = [
    "Se implementó exitosamente un sistema de seguimiento de incidentes utilizando GitHub Issues con labels personalizados, milestones y 5 issues de ejemplo, demostrando el uso de herramientas nativas de GitHub para la gestión de incidencias.",
    "La integración continua con Lighthouse CI permitió mantener estándares de calidad en rendimiento, accesibilidad, buenas prácticas y SEO, garantizando una experiencia de usuario óptima.",
    "La entrega continua mediante GitHub Actions automatizó el despliegue del frontend, reduciendo errores manuales y acelerando la publicación de nuevas características y correcciones.",
    "El despliegue en GitHub Pages y Render demostró la viabilidad de usar plataformas cloud gratuitas para aplicaciones web completas con frontend estático y backend dinámico.",
    "El proyecto integra de manera cohesiva todas las herramientas de desarrollo modernas: control de versiones (Git), CI/CD (GitHub Actions), hosting cloud (GitHub Pages, Render) y aseguramiento de calidad (Lighthouse CI).",
]
for c in conclusions:
    add_bullet(c)

# ══════════════════════════════════════════════════════
# RECOMENDACIONES
# ══════════════════════════════════════════════════════
add_heading_apa("Recomendaciones", 1)

add_heading_apa("Para el Desarrollo Futuro", 2)
recommendations_dev = [
    "Implementar autenticación de usuarios con JWT para proteger los endpoints de la API y personalizar la experiencia por usuario.",
    "Migrar la base de datos de SQLite a PostgreSQL para soportar múltiples conexiones concurrentes y mejor escalabilidad.",
    "Agregar pruebas unitarias y de integración automatizadas con pytest para garantizar la robustez del backend.",
    "Implementar un sistema de notificaciones en tiempo real utilizando WebSockets para alertar sobre nuevos incidentes.",
    "Explorar el uso de GitHub Projects para organizar los issues en tableros Kanban y mejorar la gestión visual del flujo de trabajo.",
]
for r in recommendations_dev:
    add_bullet(r)

add_heading_apa("Para el Proceso de CI/CD", 2)
recommendations_cicd = [
    "Configurar notificaciones por correo electrónico o Slack ante fallos en los pipelines de CI/CD.",
    "Implementar análisis de seguridad con herramientas como Snyk o Dependabot para detectar vulnerabilidades.",
    "Agregar un ambiente de staging (pre-producción) para validar cambios antes del despliegue a producción.",
    "Configurar pruebas de carga con herramientas como k6 o Locust para verificar el rendimiento bajo estrés.",
    "Implementar versionado semántico (SemVer) y generar changelogs automáticos con GitHub Releases.",
]
for r in recommendations_cicd:
    add_bullet(r)

add_heading_apa("Para la Documentación", 2)
recommendations_doc = [
    "Mantener un wiki vivo en el repositorio con guías de instalación, configuración y contribución.",
    "Documentar el uso de GitHub Issues con guías sobre labels, milestones y buenas prácticas.",
    "Agregar diagramas de arquitectura y flujo de datos para facilitar la comprensión del sistema.",
    "Incluir capturas de pantalla actualizadas en cada release para evidenciar el estado del proyecto.",
]
for r in recommendations_doc:
    add_bullet(r)

add_page_break()

# ══════════════════════════════════════════════════════
# REFERENCIAS
# ══════════════════════════════════════════════════════
add_page_break()
add_heading_apa("Referencias", 1)

references = [
    "American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.). https://doi.org/10.1037/0000165-000",
    "Flask. (2024). Flask Documentation (3.1.x). https://flask.palletsprojects.com/",
    "GitHub. (2024). GitHub Issues Documentation. https://docs.github.com/en/issues",
    "GitHub. (2024). GitHub Actions Documentation. https://docs.github.com/en/actions",
    "GitHub. (2024). GitHub Pages Documentation. https://docs.github.com/en/pages",
    "Google. (2024). Lighthouse CI Overview. https://github.com/GoogleChrome/lighthouse-ci",
    "Google. (2024). Lighthouse Scoring Guide. https://developer.chrome.com/docs/lighthouse/performance/performance-scoring/",
    "Gunicorn. (2024). Gunicorn WSGI Server Documentation. https://docs.gunicorn.org/",
    "MDN Web Docs. (2024). RESTful API Design. https://developer.mozilla.org/en-US/docs/Glossary/REST",
    "Pallets Project. (2024). Werkzeug Documentation. https://werkzeug.palletsprojects.com/",
    "Render. (2024). Render Documentation. https://render.com/docs",
    "SQLite. (2024). SQLite Documentation. https://www.sqlite.org/docs.html",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ─── Add page numbers ───
add_page_numbers()

# ─── Save ───
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Documentacion_Proyecto_AutoMarket_Peru.docx")
doc.save(output_path)
print(f"Documento generado exitosamente: {output_path}")
