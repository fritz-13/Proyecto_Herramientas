from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ───────────────────────────────────────────
#  PIE DE PAGINA para todas las secciones
# ───────────────────────────────────────────
def agregar_pie_pagina(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Proyecto: AutoMarket Per\u00fa | Herramientas de Desarrollo | Contenedores Docker | Julio 2026')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

# ───────────────────────────────────────────
#  INDICE
# ───────────────────────────────────────────
def agregar_indice(doc):
    doc.add_heading('\u00cdndice', level=0)
    doc.add_paragraph()

    secciones = [
        ('1.', 'Qu\u00e9 es un Contenedor'),
        ('2.', 'Beneficios de los Contenedores'),
        ('3.', 'Arquitectura de Contenedores en AutoMarket Per\u00fa'),
        ('3.1', 'Contenedor Frontend (Nginx)'),
        ('3.2', 'Contenedor Backend (Python/Flask)'),
        ('3.3', 'Orquestaci\u00f3n (Docker Compose)'),
        ('3.4', 'Proxy Reverso (Nginx)'),
        ('4.', 'Cambios Realizados'),
        ('4.1', 'Correcci\u00f3n del docker-compose.yml'),
        ('4.2', 'Optimizaci\u00f3n del Dockerfile (Frontend)'),
        ('4.3', 'Correcci\u00f3n de puertos y acceso'),
        ('4.4', 'Correcci\u00f3n del .dockerignore'),
        ('5.', 'C\u00f3mo Ejecutar el Proyecto'),
        ('6.', 'Estructura de Archivos Docker'),
        ('7.', 'Flujo de Datos'),
        ('8.', 'Soluci\u00f3n de Problemas Comunes'),
    ]

    for num, titulo in secciones:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {titulo}')
        run.font.size = Pt(11)
        if '.' not in num.strip('.'):
            run.bold = True

    doc.add_page_break()

# ───────────────────────────────────────────
#  CARATULA (extraida de avanze3-herramienta.docx)
# ───────────────────────────────────────────

# No queremos pie de pagina en la caratula
section0 = doc.sections[0]
section0.different_first_page_header_footer = True
# Hacemos que el footer de la primera pagina (caratula) este vacio
footer_first = section0.first_page_footer
footer_first.is_linked_to_previous = False

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u201cA\u00f1o de la Esperanza y el Fortalecimiento de la Democracia\u201d')
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Universidad Tecnol\u00f3gica del Per\u00fa')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ASIGNATURA:')
run.bold = True
run.font.size = Pt(12)
p.add_run('\n').font.size = Pt(12)
run2 = p.add_run('HERRAMIENTAS DE DESARROLLO')
run2.font.size = Pt(12)
run2.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DOCENTE:')
run.bold = True
run.font.size = Pt(12)
p.add_run('\n').font.size = Pt(12)
run2 = p.add_run('SOLIS FLORES ARTURO')
run2.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ESTUDIANTES:')
run.bold = True
run.font.size = Pt(12)

estudiantes = [
    ('Huaman Salas Jack', 'U22300101'),
    ('Pe\u00f1a Quispe Fritz', 'U23300415'),
    ('Tacsa Flores Jerson', 'U23273272'),
    ('Valdivia Carrasco Jos\u00e9 Daniel', 'U23273276'),
]

for nombre, codigo in estudiantes:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{nombre}\t\t{codigo}')
    run.font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Huancayo, 3 de Julio de 2026')
run.font.size = Pt(11)

doc.add_page_break()

# ───────────────────────────────────────────
#  INDICE
# ───────────────────────────────────────────
agregar_indice(doc)

# ───────────────────────────────────────────
#  TITULO DEL DOCUMENTO
# ───────────────────────────────────────────
title = doc.add_heading('Documentaci\u00f3n: Contenedores Docker en AutoMarket Per\u00fa', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Proyecto AutoMarket Per\u00fa - Contenedores con Docker')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 92, 184)

doc.add_paragraph()

# 1. Que es un Contenedor
doc.add_heading('1. Qu\u00e9 es un Contenedor', level=1)
doc.add_paragraph(
    'Un contenedor es una unidad estandarizada de software que agrupa el c\u00f3digo '
    'y todas sus dependencias para que la aplicaci\u00f3n se ejecute de manera r\u00e1pida '
    'y confiable en cualquier entorno inform\u00e1tico.'
)
doc.add_paragraph(
    'A diferencia de una m\u00e1quina virtual, un contenedor comparte el kernel del '
    'sistema operativo host, lo que lo hace m\u00e1s ligero y r\u00e1pido de iniciar.'
)

# 2. Beneficios
doc.add_heading('2. Beneficios de los Contenedores', level=1)

beneficios = [
    ('Portabilidad', 'La aplicaci\u00f3n funciona igual en cualquier m\u00e1quina con Docker instalado'),
    ('Consistencia', 'Elimina el problema \u201cen mi m\u00e1quina funciona\u201d'),
    ('Aislamiento', 'Cada servicio corre en su propio contenedor sin afectar a otros'),
    ('Escalabilidad', 'Facilita duplicar o escalar servicios seg\u00fan la demanda'),
    ('CI/CD', 'Compatible con GitHub Actions, Kubernetes y cloud providers (AWS, Azure, GCP)')
]

for titulo, desc in beneficios:
    p = doc.add_paragraph()
    run = p.add_run(f'{titulo}: ')
    run.bold = True
    p.add_run(desc)

# 3. Arquitectura
doc.add_heading('3. Arquitectura de Contenedores en AutoMarket Per\u00fa', level=1)
doc.add_paragraph(
    'El proyecto utiliza Docker Compose para orquestar dos contenedores:'
)

table = doc.add_table(rows=1, cols=2)
table.style = 'Medium Shading 1 Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'FRONTEND (Nginx)'
hdr_cells[1].text = 'BACKEND (Python/Flask)'

row = table.add_row().cells
row[0].text = 'Puerto: 80'
row[1].text = 'Puerto: 5000'

row = table.add_row().cells
row[0].text = 'index.html, styles.css, navbar/, main/, Login-register/'
row[1].text = 'app.py, wsgi.py, requirements.txt'

row = table.add_row().cells
row[0].text = 'Sirve archivos est\u00e1ticos y proxy reverso'
row[1].text = 'API REST con Flask y Gunicorn'

doc.add_paragraph()

doc.add_heading('3.1 Contenedor Frontend (Nginx)', level=2)
doc.add_paragraph('Archivo: Dockerfile (ra\u00edz del proyecto)')

items = [
    'Imagen base: nginx:alpine (ligera y optimizada)',
    'Funci\u00f3n: Sirve los archivos est\u00e1ticos (HTML, CSS, JS)',
    'Puerto: 80 (accesible desde http://localhost)',
    'Configuraci\u00f3n: nginx.conf act\u00faa como proxy reverso',
    'Copia todo el proyecto con COPY . /usr/share/nginx/html/'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Contenedor Backend (Python/Flask)', level=2)
doc.add_paragraph('Archivo: backend/Dockerfile')

items = [
    'Imagen base: python:3.10-slim',
    'Funci\u00f3n: Ejecuta la API REST con Flask y Gunicorn',
    'Puerto: 5000 (accesible desde http://localhost:5000)',
    'Dependencias: flask, flask-cors, gunicorn'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 Orquestaci\u00f3n (Docker Compose)', level=2)
doc.add_paragraph('Archivo: docker-compose.yml')

items = [
    'Define ambos servicios (frontend y backend)',
    'Configura puertos y dependencias',
    'Crea un volumen persistente para datos (db-data)',
    'Usa restart: unless-stopped para alta disponibilidad'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.4 Proxy Reverso (Nginx)', level=2)
doc.add_paragraph('Archivo: nginx.conf')

items = [
    'Escucha en puerto 80',
    'Redirige solicitudes /api/* al backend en puerto 5000',
    'Sirve archivos est\u00e1ticos del frontend',
    'Configuraci\u00f3n de cabeceras HTTP para proxy'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# 4. Cambios Realizados
doc.add_heading('4. Cambios Realizados', level=1)

doc.add_heading('4.1 Correcci\u00f3n del docker-compose.yml', level=2)
doc.add_paragraph(
    'Se elimin\u00f3 la l\u00ednea version: \"3.8\" porque es un atributo obsoleto en '
    'versiones recientes de Docker Compose. Docker Compose v2+ no requiere esta especificaci\u00f3n.'
)

p = doc.add_paragraph()
run = p.add_run('Antes:')
run.bold = True
doc.add_paragraph('version: "3.8"\n\nservices:\n  frontend:\n    ...')

p = doc.add_paragraph()
run = p.add_run('Despu\u00e9s:')
run.bold = True
doc.add_paragraph('services:\n  frontend:\n    ...')

doc.add_heading('4.2 Optimizaci\u00f3n del Dockerfile (Frontend)', level=2)
doc.add_paragraph(
    'Se simplific\u00f3 el Dockerfile del frontend reemplazando m\u00faltiples '
    'instrucciones COPY por una sola instrucci\u00f3n COPY . /usr/share/nginx/html/. '
    'Esto asegura que todos los archivos del proyecto se copien al contenedor '
    'sin riesgo de omitir archivos necesarios.'
)

p = doc.add_paragraph()
run = p.add_run('Antes:')
run.bold = True
doc.add_paragraph('COPY index.html /usr/share/nginx/html/\nCOPY styles.css /usr/share/nginx/html/\nCOPY footer.css /usr/share/nginx/html/\nCOPY incidentes.html /usr/share/nginx/html/\n...')

p = doc.add_paragraph()
run = p.add_run('Despu\u00e9s:')
run.bold = True
doc.add_paragraph('COPY . /usr/share/nginx/html/')

doc.add_heading('4.3 Correcci\u00f3n de puertos y acceso', level=2)
doc.add_paragraph(
    'Se identific\u00f3 que la API Flask no tiene ruta en la ra\u00edz (/), '
    'por lo que las URLs correctas de acceso son:'
)

urls = [
    'Frontend: http://localhost',
    'Backend API directo: http://localhost:5000/api/incidentes',
    'Backend v\u00eda proxy Nginx: http://localhost/api/incidentes',
    'Ejecutar en segundo plano: docker-compose up -d'
]
for u in urls:
    doc.add_paragraph(u, style='List Bullet')

doc.add_heading('4.4 Correcci\u00f3n del .dockerignore', level=2)
doc.add_paragraph(
    'Se agregaron entradas al .dockerignore para excluir archivos innecesarios '
    'de la construcci\u00f3n de la imagen Docker, reduciendo el tama\u00f1o del contexto '
    'de build y mejorando los tiempos de construcci\u00f3n.'
)

p = doc.add_paragraph()
run = p.add_run('Archivos ignorados en la construcci\u00f3n Docker:')
run.bold = True
doc.add_paragraph('.git, .github, .lighthouseci, *.md, *.docx, *.pptx, render.yaml, lighthouse-report.html, Roadmap, opencode, cd')

# 5. Como Ejecutar
doc.add_heading('5. C\u00f3mo Ejecutar el Proyecto', level=1)

doc.add_heading('Prerrequisitos', level=2)
doc.add_paragraph('Docker Desktop instalado', style='List Bullet')
doc.add_paragraph('WSL2 habilitado con distribuci\u00f3n Ubuntu', style='List Bullet')

doc.add_heading('Pasos', level=2)
doc.add_paragraph('1. Abrir PowerShell en la carpeta del proyecto', style='List Number')
doc.add_paragraph('2. Ejecutar: docker-compose up --build', style='List Number')
doc.add_paragraph('3. Opcional: usar -d para ejecutar en segundo plano: docker-compose up -d --build', style='List Number')

p = doc.add_paragraph()
run = p.add_run('Acceder desde el navegador:')
run.bold = True
doc.add_paragraph('Frontend: http://localhost', style='List Bullet')
doc.add_paragraph('Backend API: http://localhost:5000/api/incidentes', style='List Bullet')
doc.add_paragraph('API v\u00eda proxy: http://localhost/api/incidentes', style='List Bullet')

doc.add_heading('Comandos \u00datiles', level=2)

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Medium Shading 1 Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Comando'
hdr_cells[1].text = 'Descripci\u00f3n'

comandos = [
    ('docker-compose up --build', 'Construir e iniciar contenedores'),
    ('docker-compose up -d', 'Iniciar en segundo plano'),
    ('docker-compose down', 'Detener contenedores'),
    ('docker-compose ps', 'Ver contenedores en ejecuci\u00f3n'),
    ('docker-compose logs', 'Ver logs de los contenedores'),
    ('docker-compose exec <servicio> <cmd>', 'Ejecutar comando en un contenedor'),
    ('docker-compose restart', 'Reiniciar contenedores'),
]

for cmd, desc in comandos:
    row = table2.add_row().cells
    row[0].text = cmd
    row[1].text = desc

doc.add_paragraph()

# 6. Estructura de Archivos
doc.add_heading('6. Estructura de Archivos Docker', level=1)

estructura = """Proyecto_Herramientas/
+-- Dockerfile              # Contenedor Frontend (Nginx)
+-- docker-compose.yml      # Orquestador de contenedores
+-- nginx.conf              # Configuracion Nginx (proxy reverso)
+-- .dockerignore           # Archivos ignorados en build
+-- index.html              # Pagina principal
+-- styles.css              # Estilos generales
+-- footer.css              # Estilos del footer
+-- incidentes.html         # Dashboard de incidentes
+-- incidentes.css          # Estilos del dashboard
+-- navbar/                 # Recursos del navbar
+-- main/                   # Recursos del main
+-- Login-register/         # Paginas de login y registro
+-- backend/
|   +-- Dockerfile          # Contenedor Backend (Python)
|   +-- requirements.txt    # Dependencias Python
|   +-- app.py              # Aplicacion Flask (API REST)
|   +-- wsgi.py             # Punto de entrada WSGI
|   +-- incidentes.db       # Base de datos SQLite
+-- ..."""

p = doc.add_paragraph()
run = p.add_run(estructura)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# 7. Flujo de Datos
doc.add_heading('7. Flujo de Datos', level=1)

flujo = """Usuario (Navegador)
       |
       v
   Puerto 80 (Nginx)
       |
       +--> Archivos Estaticos (HTML/CSS/JS)
       |
       +--> /api/* --> Puerto 5000 (Flask Backend)
                              |
                              v
                      Base de Datos (incidentes.db)"""

p = doc.add_paragraph()
run = p.add_run(flujo)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# 8. Solucion de Problemas
doc.add_heading('8. Soluci\u00f3n de Problemas Comunes', level=1)

doc.add_heading('8.1 Error 404 en el navegador', level=2)
doc.add_paragraph(
    'Si http://localhost muestra un error 404, verificar que los contenedores '
    'est\u00e9n en ejecuci\u00f3n con docker-compose ps y que el puerto 80 '
    'no est\u00e9 siendo usado por otro programa en el host.'
)

doc.add_heading('8.2 El backend no responde en localhost:5000', level=2)
doc.add_paragraph(
    'La API Flask no tiene ruta en la ra\u00edz (/). Usar http://localhost:5000/api/incidentes '
    'para listar los incidentes. Si no responde, revisar los logs con docker-compose logs backend.'
)

doc.add_heading('8.3 Puerto 80 ocupado en Windows', level=2)
doc.add_paragraph(
    'Si otro programa ya est\u00e1 usando el puerto 80, cambiar el mapeo en '
    'docker-compose.yml a un puerto alternativo (ej: \"8080:80\") y acceder '
    'a http://localhost:8080.'
)

doc.add_paragraph()
doc.add_paragraph()

# ───────────────────────────────────────────
#  PIE DE PAGINA (a partir de la segunda pagina)
# ───────────────────────────────────────────
for section in doc.sections:
    section.different_first_page_header_footer = True
    agregar_pie_pagina(section)

doc.save(r'C:\Users\fritz\OneDrive\Desktop\UTP\HERRAMIENTA DE DESARROLLO\avanze 2\Proyecto_Herramientas\Documentacion_Contenedores.docx')
print('Documento Word generado correctamente.')
